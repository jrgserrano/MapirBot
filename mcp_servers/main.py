from mcp.server.fastmcp import FastMCP
import chromadb
from langchain_text_splitters import RecursiveCharacterTextSplitter, TokenTextSplitter
from graphiti_core import Graphiti
from graphiti_core.llm_client import OpenAIClient, LLMConfig
from graphiti_core.embedder import OpenAIEmbedder, OpenAIEmbedderConfig
from graphiti_core.nodes import EpisodeType
from datetime import datetime, timezone
import os
import base64
import io
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv

load_dotenv()

mcp = FastMCP("mapir_mcp")

chroma_client = chromadb.PersistentClient(path="/Users/jorseme/Desktop/MapirBot/V2/mcp_server/chroma_db")
collection = chroma_client.get_or_create_collection("mapir_docs")

neo4j_uri = os.environ.get('NEO4J_URI', 'bolt://localhost:7687')
neo4j_user = os.environ.get('NEO4J_USER', 'neo4j')
neo4j_password = os.environ.get('NEO4J_PASSWORD', 'neo4j_changed')

llm_name = 'google/gemma-3-4b'
openai_base_url = 'http://localhost:1234/v1'

llm_config = LLMConfig(
    api_key="lm-studio",
    base_url=openai_base_url,
    model=llm_name,
    small_model=llm_name
)
llm_client = OpenAIClient(config=llm_config)

embedder_config = OpenAIEmbedderConfig(
    api_key="ollama",
    base_url='http://localhost:11434/v1',
    embedding_model='mxbai-embed-large'
)
embedder = OpenAIEmbedder(config=embedder_config)

graph = Graphiti(
    uri=neo4j_uri,
    user=neo4j_user,
    password=neo4j_password,
    llm_client=llm_client,
    embedder=embedder
)

@mcp.tool()
async def search_rag(query: str, n_results: int = 3) -> str:
    """Search for relevant information in ChromaDB."""
    results = collection.query(
        query_texts=[query],
        n_results=n_results
    )
    return "\n".join(results["documents"][0])

@mcp.tool()
async def search_graph(query: str, n_results: int = 3) -> str:
    """Search for relevant information in the knowledge graph."""
    results = graph.search(query, n_results)
    return "\n".join([f"- {edge.fact}" for edge in results])

@mcp.tool()
async def hybrid_search(query: str, n_results: int = 3) -> str:
    """Search for relevant information in the knowledge base."""
    rag_results = await search_rag(query, n_results)
    graph_results = await search_graph(query, n_results)
    return f"Vector Context:\n{rag_results}\n\nGraph Context:\n{graph_results}"

@mcp.tool()
async def add_file(file_name: str, file_data_base64: str = None, local_path: str = None) -> str:
    """Add a new file to the knowledge base."""

    content = ""

    if file_data_base64:
        file_data = io.BytesIO(base64.b64decode(file_data_base64))
    elif local_path:
        with open(local_path, "rb") as f:
            file_data = io.BytesIO(f.read())
    else:
        return "No file data provided."

    ext = file_name.split(".")[-1]

    if ext == "pdf":
        reader = PdfReader(file_data)
        content = "\n".join([page.extract_text() for page in reader.pages])
    elif ext == "docx":
        doc = Document(file_data)
        content = "\n".join([para.text for para in doc.paragraphs])
    else:
        content = file_data.getvalue().decode("utf-8")

    if content:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,
            chunk_overlap=100,
            add_start_index=True,
            separators=["\n\n", "\n", ".", " "],
        )
        chunks = text_splitter.split_text(content)
        ids = [f"{file_name}_{i}" for i in range(len(chunks))]

        metadatas = [{"source": file_name}]*len(chunks)
    
        collection.add(
            documents=chunks,
            ids=ids,
            metadatas=metadatas
        )

        await graph.add_block(content, source=file_name)

        return f"Document '{file_name}' added successfully. {len(chunks)} chunks added."
    
    return "No content found in the file."

def extract_text_from_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"No se encuentra el archivo en: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    content = ""
    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            content = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif ext == ".docx":
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext in [".txt", ".md", ".cpp", ".hpp", ".py"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        return content.strip()
    except Exception as e:
        raise Exception(f"Error procesando {os.path.basename(file_path)}: {str(e)}")


@mcp.tool()
async def process_and_summarize_doc(file_path: str, category: str = "robotica"):
    """
    Ingesta archivos dividiéndolos en bloques para evitar errores de contexto y timeouts.
    """
    # 1. Extraer texto plano
    content = extract_text_from_file(file_path)
    name = os.path.basename(file_path)

    # 2. Configurar el fraccionamiento (Chunking)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2500,
        chunk_overlap=250,
        add_start_index=True
    )
    chunks = text_splitter.split_text(content)

    # 3. Guardar en ChromaDB (RAG Vectorial)
    # Esto es rápido y no depende del LLM
    collection.add(
        documents=chunks,
        ids=[f"{name}_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_path, "is_summary": False}] * len(chunks)
    )


    # 4. Sincronizar con Graphiti (Neo4j) - PROCESO POR BLOQUES
    # Procesamos solo los primeros 5 bloques para asegurar estabilidad
    # Cada bloque toma unos segundos en procesarse (Prompt Processing)
    max_graph_chunks = 5 
    print(f"Iniciando construcción de grafo para {name}...")
    
    for i, chunk in enumerate(chunks[:max_graph_chunks]):
        try:
            print(f"Procesando bloque {i+1}/{max_graph_chunks}...")
            await graph.add_episode(
                name=f"{name}_p{i}", 
                episode_body=chunk,
                source_description=f"Sección {i+1} de {name}",
                source=EpisodeType.text,
                reference_time=datetime.now(timezone.utc),
            )
        except Exception as e:
            print(f"Error en bloque {i}: {e}")
            continue # Intentamos con el siguiente si uno falla

    return {
        "status": "success",
        "message": f"Archivo '{name}' procesado.",
        "info": f"Vectores creados: {len(chunks)}. Bloques en grafo: {min(len(chunks), max_graph_chunks)}.",
        "resumen_ia": resumen_clave
    }

if __name__ == "__main__":
    mcp.run()
    